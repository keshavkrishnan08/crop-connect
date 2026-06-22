#!/usr/bin/env python3
"""Build the investor pre-read as a clean, editable .docx (memo format)."""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BRAND = RGBColor(0x23, 0x5C, 0x3A)
INK = RGBColor(0x16, 0x24, 0x1C)
GREY = RGBColor(0x5b, 0x6b, 0x60)

doc = Document()
s = doc.sections[0]
s.top_margin = s.bottom_margin = Inches(0.8)
s.left_margin = s.right_margin = Inches(0.9)
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

def runs(p, text, size=10.5, color=INK):
    for idx, seg in enumerate(text.split("**")):
        bold = idx % 2 == 1
        for part in re.split(r"(\{[^}]*\})", seg):
            if not part:
                continue
            if part[0] == "{" and part[-1] == "}":
                r = p.add_run(part[1:-1]); r.font.superscript = True; r.font.size = Pt(7); r.font.color.rgb = BRAND
            else:
                r = p.add_run(part); r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = color

def title(text):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(19); r.font.color.rgb = INK

def heading(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(11); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = BRAND

def para(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; runs(p, text)

def bullet(text):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3); runs(p, text)

def numbered(text):
    p = doc.add_paragraph(style="List Number"); p.paragraph_format.space_after = Pt(4); runs(p, text)

def subbullet(text):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.65); runs(p, text, size=10)

def note(text):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4); runs(p, text, size=9.5, color=GREY)

def sources(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); runs(p, text, size=8.5, color=GREY)

# ---- content (memo format) ----
title("CropConnect Investor Pre-Read")

heading("Summary")
para("CropConnect is an AI agent that does one thing for a restaurant: it runs local sourcing end to end, finding the farms, signing the contracts, scheduling weekly delivery, holding payment in escrow, and tracking the added margin. The owner names an ingredient and puts the dish on the menu; the software does the rest, automatically. The business is **asset-light**: we own no warehouses, trucks, or inventory, the food is billed at the farm's price, and delivery is third-party, so our only revenue is a **monthly service fee** with software-level margins. The product is **built and live, developed by the founder solo**. We are pre-revenue, raising **$50,000 to $100,000** on a SAFE to land the first **15 paying kitchens in Indianapolis**.")

heading("The Opportunity")
para("Until now, running local sourcing for a restaurant required a salaried human coordinator. An AI agent can now do that work directly: it **connects a restaurant to the right local farms, drafts and negotiates custom supply contracts, and runs the ongoing services on its own, including delivery tracking, escrow payments, and reconciliation.** Software can now perform the coordination that used to require a person, which is what makes offering local sourcing profitable. That capability is recent, and it is the opportunity.")

heading("Market and Timing")
para("Independent full-service restaurants net about **3% to 5% of revenue**, with food and labor consuming roughly **60 cents of every dollar**.{1} Most growth has to come from charging more for the same seats, and local sourcing is a proven way to do it: **38% of U.S. diners** are more likely to choose a restaurant offering locally sourced food,{2} and the share rating ingredient transparency as important rose from **69% in 2018 to 76% in 2023**.{3} Sourcing locally by hand requires finding farms, vetting them, negotiating, contracting, scheduling, and covering shortfalls.")
para("The U.S. restaurant market is about **$1.1 trillion** across roughly **1 million** locations.{4} Of those, we estimate about **150,000 are independent full-service restaurants that would plausibly buy**: chef-driven, mid-to-upscale operators that value local sourcing. That is a serviceable market of roughly **$1.6 billion a year** at our pricing.{5} In the **Indianapolis** metro, of about 4,500 restaurants we estimate roughly **1,500 fit that profile**, about **$16 million a year**.{5} We start in Indianapolis and repeat the same playbook city by city; each metro is a self-contained, profitable unit. Beyond restaurants, the farm relationships and the allocation layer become a demand aggregator for growers, a **farming platform** larger than restaurant sourcing alone.")
sources("Sources.  1. Restaurant net-margin and prime-cost benchmarks: National Restaurant Association; Restaurant365.  2. National Restaurant Association, 2022 State of the Restaurant Industry.  3. FMI, transparency and foodservice trends, 2023.  4. National Restaurant Association, 2024 industry forecast.  5. U.S. and Indianapolis target-segment counts and the revenue estimates derived from them are internal estimates.")

heading("Product")
para("The restaurant enters one need, say 40 lbs of heirloom tomatoes a week, and the agent runs the entire workflow **automatically**: it ranks vetted local farms by crop, distance, reliability, and price; splits the order across farms if one cannot cover it; drafts a contract and negotiates quality terms; schedules weekly third-party delivery; holds payment in escrow until delivery is confirmed; and tracks the added margin per dish. The founder built it as a working product, live today; each step runs in software without manual intervention. A restaurant can run a free instant audit before signing up.")

heading("Business Model")
para("Revenue comes only from a monthly service fee. Food is billed at the farm's price and delivery is passed through at third-party cost, both with no markup. The fee is a **$299** base plus a per-item fee that scales with volume, blending to about **$896 per restaurant per month** (about $10,700/year). Because our costs are software, agent compute, and a thin coordination layer, the service-fee gross margin is high.")

heading("Unit Economics")
para("These are targets, not results; we are pre-revenue and will validate them against the first cohort.")
bullet("**ARPA:** about $896/month at maturity, starting near $400 on the first ingredient and growing as items are added.")
bullet("**Cost to serve one restaurant: about $130/month**, for a gross margin near **85%** (about $766/month of gross profit). The components:")
subbullet("Agent tokens on the **Anthropic Claude API**, about $45 (farm matching, drafting and negotiating contracts, monthly audits, and specials).")
subbullet("Payments and escrow on **Stripe Connect**, about $50 (card processing on the service fee; food held in escrow moves over ACH to keep fees low).")
subbullet("Database, auth, and hosting on **Supabase and Vercel**, about $15.")
subbullet("Light human oversight, about $20.")
subbullet("Delivery runs through a **third-party courier API (for example Uber Direct)** and is billed through to the restaurant at cost, so it is not a margin cost to us.")
bullet("**CAC: about $450 per restaurant** = a sales commission of about $400, plus one week of operating cost (agent tokens and delivery, about $50) for the free first ingredient.")
bullet("**Payback:** under one month at target margin.")
bullet("**Monthly churn:** assume about 3%, reflecting normal restaurant turnover; the embedded sourcing, contracts, and deliveries should make real churn lower.")
para("We treat lifetime value as a hypothesis to test with data, not a number to lead with.")

heading("Go-to-Market")
para("We win one metro at a time and earn density before expanding. The Indianapolis plan, in order:")
numbered("**Onboard the farms.** Pre-sign 8 to 12 vetted local farms across the core crops so supply and fill are guaranteed before we sell a single restaurant. This removes the cold-start problem.")
numbered("**Show restaurants their savings.** For each target we run the free audit on their public menu and bring a specific, quantified estimate of the money they could save and the margin they are leaving behind.")
numbered("**Win them with a risk-free offer.** First ingredient free, free ten-minute setup, cancel any month, backed by a guarantee. We start with one ingredient (about $400/month) and grow the account toward the $896 blended ARPA.")
numbered("**Expand metro by metro.** At roughly 30 to 40 kitchens and 15 or more farms in Indianapolis, open the next city with the same playbook.")

heading("Why We Win")
para("Investors have put real money into food-supply software, which shows the demand is real. Misfits Market reached about a **$2B** valuation, Choco about **$1.2B**, and Afresh has raised about **$148M** for grocery forecasting.{6} None offers done-for-you local sourcing for restaurants. Choco digitizes ordering between a restaurant and the distributors it already uses, but it does not find or contract local farms. Afresh forecasts inventory for grocery chains. Misfits Market sells surplus produce to consumers. CropConnect is the only one that runs local sourcing end to end for restaurants: finding the farms, contracting them, and managing delivery, escrow, and margin.")
para("The companies that tried to deliver local food themselves struggled because they owned the perishable inventory and the logistics. We stay asset-light and pass the food through at cost, which is why the economics work for us. The advantage also compounds: vetting a city's farms once forces a rival to rebuild the entire local supply base, every contract and delivery trains the agent on local supply and demand, and a kitchen that runs its sourcing, contracts, and deliveries through us would have to rebuild its supply chain by hand to leave.")
sources("Source.  6. Valuations are peak reported figures from press and PitchBook: Misfits Market (about $2B, 2021) and Choco (about $1.2B, 2022); Afresh has raised about $148M.")

heading("Team")
para("**Keshav Krishnan**, founder. His research on the **economic cost of climate-driven migration** gave him a data-level view of how fragile long-distance food supply chains are and why regional, resilient sourcing is where the system is heading, the conviction behind CropConnect. He also **built the entire product solo**, from the agent to the contracts and escrow, which is why a working platform already exists at the pre-seed stage. The first hire this round funds is a **Head of Sourcing**, who builds and manages the local farm supply and uses existing restaurant relationships to open doors. The strongest candidates are **produce-distributor sales reps or territory managers** (for example at Indianapolis Fruit, Piazza Produce, or What Chefs Want) and **food-hub or farm-cooperative managers** (for example Hoosier Harvest Market or the Farmers Cooperative Food Hub), who already carry a book of restaurant accounts and know the local growers. A modest base plus meaningful equity is attractive to a strong rep ready to trade a commission grind for ownership.")
note("To complete before sending: one concrete figure from your climate-migration research, and the name of a sourcing hire if you have one lined up.")

heading("Status and Ask")
para("The product is built and live; we are **pre-revenue**. We are raising **$50,000 to $100,000** on a post-money SAFE (**$4M** cap, **20%** discount, MFN), with about **80%** going to sales. The round funds one milestone: **15 paying kitchens and about $13,000 in monthly recurring revenue in Indianapolis**, the proof point to raise a priced seed.")

doc.save("pitch/cropconnect-preread.docx")
print("saved pitch/cropconnect-preread.docx")
